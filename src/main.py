from pylibdmtx.pylibdmtx import encode
from PIL import Image
import docx
from docx.shared import Inches, Cm, Mm


def add_dmcs(doc, batch_name, codes_amount, production_date, supplier_code, part_number, size):
    docu.add_heading(f'{batch_name}: {supplier_code}', 3)

    table = doc.add_table(rows=1,cols=4)

    for x in range(1,codes_amount+1):
        if x < 10:
            serial_number = f'100000{x}'
        else:
            serial_number = f'10000{x}'

        code = production_date + chr(29) + supplier_code + chr(29) + serial_number + chr(29) + part_number + chr(29)
        

        encoded = encode(code.encode('utf8'))
        img = Image.frombytes('RGB', (encoded.width, encoded.height), encoded.pixels)
        img.save(f'{serial_number}.png')

        pic_cel = table.rows[0].cells[x % 4].add_paragraph()
        run = pic_cel.add_run()

        run.add_picture(f'{serial_number}.png', width=Cm(size), height=Cm(size))


print("done")


docu = docx.Document()
docu.add_heading('Magna Bari DMCs', 0)
add_dmcs(docu, 'carrier plate', 15, '231009', '88888822', 'M3333331_01', 4)
add_dmcs(docu, 'carrier plate', 15, '231009', '88888833', 'M3333331_01', 4)
add_dmcs(docu, 'valve plate', 15, '231009', '77777722', 'M4444441_01', 4)
add_dmcs(docu, 'valve plate', 15, '231009', '77777733', 'M4444441_01', 4)
add_dmcs(docu, 'intermediate plate', 15, '231009', '66666633', 'M5555551_01', 4)
add_dmcs(docu, 'intermediate plate', 15, '231009', '66666644', 'M5555551_01', 4)
docu.save('MagnaBariDMC.docx')







# print(decode(Image.open('dmtx.png')))

